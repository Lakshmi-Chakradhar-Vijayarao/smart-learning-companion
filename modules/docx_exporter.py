from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def export_docx(sections, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for section in sections:
        title = section.get("title", "").strip()
        content = section.get("content", "").strip()

        # Add title as Heading 1
        if title:
            title_para = doc.add_paragraph(title)
            title_para.style = doc.styles['Heading 1']
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add content as paragraph(s)
        if content:
            for para in content.split("\n\n"):
                doc.add_paragraph(para.strip())

    doc.save(output_path)
